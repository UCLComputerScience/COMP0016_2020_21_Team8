# -*- coding: utf-8 -*-
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.layout import *
from pdfminer.converter import PDFPageAggregator
import pdftitle
import docx

# def getData(dict, type):
#     if type in dict:
#         value = dict.get(type)
#         if isinstance(value, PDFObjRef):
#             value = resolve1(value)
#         try:
#             value = value.decode(chardet.detect(value).get('encoding'))
#         except:
#             value = 'n/a'
#     else:
#         value = 'n/a'

#     if value == '':
#         value = 'n/a'

#     return value


def readPdf(pdfpath, txtpath):
    fp = open(pdfpath, 'rb')
    # 来创建一个pdf文档分析器
    parser = PDFParser(fp)
    # 创建一个PDF文档对象存储文档结构
    document = PDFDocument(parser)
    # 检查文件是否允许文本提取
    if not document.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        # 创建一个PDF资源管理器对象来存储共赏资源
        rsrcmgr = PDFResourceManager()
        # 设定参数进行分析
        laparams = LAParams()
        # 创建一个PDF设备对象
        # device=PDFDevice(rsrcmgr)
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        # 创建一个PDF解释器对象
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        # 处理每一页

        #dict = document.info[0]

        title = pdftitle.get_title_from_file(pdfpath)
        if title is None:
            title = 'n/a'

        # date = getData(dict,'CreationDate')

        # print(title)
        # print(date)
        # print(dict)

        with open(txtpath, 'w') as f:
            count = 0
            content = ""
            for page in PDFPage.create_pages(document):
                pageContent = ""
                interpreter.process_page(page)
                # 接受该页面的LTPage对象
                layout = device.get_result()
                for x in layout:
                    if(isinstance(x, LTTextBoxHorizontal)):
                        pageContent += x.get_text()
                content += pageContent.strip() + " "
                count += 1
            print(len(content))
            f.write(content)
        return (title, count)


def readDocx(docxpath, txtpath):
    # 获取文档对象
    document = docx.Document(docxpath)
    # 完整的text：
    docx_text = ""
    with open(txtpath, 'w') as f:
        for para in document.paragraphs:
            docx_text += para.text + '\n'
        f.write(docx_text)
    print(len(docx_text))

    return (document.core_properties.title,len(docx_text)//2700)

if __name__ == "__main__":

    print(readDocx('text-summarizationFortest/test.docx','text-summarizationFortest/test.txt'))
    print(readPdf('text-summarizationFortest/test.pdf','text-summarizationFortest/test1.txt'))
